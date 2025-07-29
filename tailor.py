import requests
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches

# Function to extract the text and structure from a word resume file
def extract_text_and_structure(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip() == "":
            full_text.append(("", "HR", para.runs, para.alignment, 1))
        else:
            for run in para.runs:
                if run.font.size:
                    font = round(run.font.size.pt, 2)
            full_text.append((para.text, para.style.name, para.runs, para.alignment, font))
    return full_text

# Function to tailor the text of the resume
def tailor_text(resume_text, job_description):
    prompt = f"""
    You are a resume tailoring assistant. Your task is to rewrite and improve the provided resume so that it better matches the job description, while keeping the structure and content grounded in the original resume only.

    Here is the resume followed by the job description:
    Resume:
    {resume_text}
    Job Description:
    {job_description}

    == FORMAT AND STRUCTURE REQUIREMENTS ==
    1. Keep all original section titles: "Professional Summary", "Education", "Experience", "Projects", and "Key Skills".
    2. For each experience or project entry, include exactly 3 to 4 bullet points, each starting with "-".
    3. In "Key Skills", maintain the 4 subcategories: Languages/Frameworks, Concepts, Tools, and Soft Skills.
    4. You may rewrite or improve any line or sentence, but do NOT:
    - Invent technologies, tools, or concepts the original resume does not already mention.
    - Reuse entire sentences from the original resume unless they already fit the job description well.
    5. Do not add any new sections or headings that are not present in the example format.
    6. Do not include any formatting like asterisks, bolding, or markdown. Just plain text and "-" for bullet points.

    Please strictly follow the structure shown below. The *content* is just filler—replace it with real tailored content from the original resume and job description. This example is only to demonstrate structure:
    == DESIRED OUTPUT EXAMPLE ==
    Professional Summary
    Tailored summary sentence that aligns with job description.

    Education
    University of Ottawa                                                         Sept. 2020 – Aug. 2025
    -BASc Software Engineering (GPA 9.18)

    Experience
    Knak – Full Stack Developer (Co-op)                                                        May 2023 – Dec. 2023
    -Bullet point one tailored to job
    -Bullet point two
    -Bullet point three
    -Bullet point four

    Solace – QA Engineer (Co-op)                                                                  Sept. 2022 – Dec. 2022
    -Bullet point one
    -Bullet point two
    -Bullet point three
    -Bullet point four

    FINTRAC – Software Developer (Co-op)                                                Feb. 2022 – Apr. 2022
    -Bullet point one
    -Bullet point two
    -Bullet point three
    -Bullet point four

    Projects
    Club Website Platform                                                                    Sept. 2024 – Present
    -Bullet point one
    -Bullet point two
    -Bullet point three
    -Bullet point four

    Key Skills
    Languages/Frameworks: React, TypeScript, Java, SQL
    Concepts: REST APIs, CI/CD, Agile, Unit Testing
    Tools: Git, Docker, Jenkins
    Soft Skills: Communication, Teamwork, Adaptability
    == END EXAMPLE ==

    Now return a tailored resume using this exact structure and only relevant information grounded in the original resume and job description.
    """
    
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "mistral:instruct", 
                "prompt": prompt,
                "stream": False,
            }
        )
        if response.status_code == 200:
                try:
                    response_data = response.json()
                    tailored_text = response_data.get("response", "")
                    return tailored_text
                except ValueError as e:
                    print(f"Error decoding JSON: {e}")
        else:
            print(f"Failed to get a successful response. Status code: {response.status_code}")
            print(f"Response content: {response.content}")
    except requests.exceptions.RequestException as e:
        print(f"Request failed: {e}")
    
    return None

# Function to create a new word file with the contents of the tailored resume
def create_tailored_docx(tailored_text, output_path):
    tailored_doc = Document()

    # Setting nomral 1 margins
    section = tailored_doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Define constants
    default_font = "Times New Roman"
    default_size = Pt(11)
    section_headers = { "Professional Summary", "Education", "Experience", "Projects", "Key Skills" }
    sub_header_prefix = { "University of Ottawa", "Knak", "Solace", "FINTRAC", "Club Website Platform" }

    # Create static information section
    para_name = tailored_doc.add_paragraph()  # Create new paragraph
    para_name.alignment = 1  # Center alignment
    para_name.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    para_name.paragraph_format.line_spacing = 1.0
    run_name = para_name.add_run('Zeyu Shao')  # Create new paragraph run
    run_name.font.name = default_font  # Setting line spacing to 1.0
    run_name.font.size = Pt(20)  # Set font size to 20
    run_name.bold = True  # Set font to bold

    para_address = tailored_doc.add_paragraph()  
    para_address.alignment = 1
    para_address.paragraph_format.space_after = Pt(0)
    para_address.paragraph_format.line_spacing = 1.0
    run_address = para_address.add_run('530 Anchor Circle, Ottawa, ON K4M 0X5')
    run_address.font.name = default_font
    run_address.font.size = Pt(12)

    para_info = tailored_doc.add_paragraph()  
    para_info.alignment = 1
    para_info.paragraph_format.space_after = Pt(0)
    para_info.paragraph_format.line_spacing = 1.0
    run_info = para_info.add_run('Phone: (343)-777-3529   E-mail: zeyu.shao@hotmail.com')
    run_info.font.name = default_font
    run_info.font.size = Pt(12)

    para_separator = tailored_doc.add_paragraph()
    para_separator.alignment = 1 
    para_separator.paragraph_format.space_after = Pt(0)
    run_seperator = para_separator.add_run('.')
    run_seperator.font.size = Pt(1)
    add_paragraph_border(para_separator)

    # Split the tailored text into paragraphs
    tailored_paragraphs = tailored_text.split('\n')

    for para_text in tailored_paragraphs:
        para_text = para_text.strip()
        if para_text == "":
            # Add a horizontal line implemented as a border
            para = tailored_doc.add_paragraph()
            para.alignment = 1
            para.paragraph_format.space_after = Pt(0)

            new_run = para.add_run('.')
            new_run.font.size = Pt(1) 

            add_paragraph_border(para)
            continue  # Skip section separators

        # Detect section titles by simple rules (adjust as needed)
        if para_text in section_headers:
            para = tailored_doc.add_paragraph()

            run = para.add_run(para_text)
            run.font.name = default_font
            run.font.size = Pt(14)
            run.bold = True
        elif para_text.startswith("-"):
            para_text = para_text[1:]

            para = tailored_doc.add_paragraph(style='List Bullet 2')
            
            run = para.add_run(para_text)
            run.font.name = default_font
            run.font.size = default_size
        elif any(para_text.startswith(word) for word in sub_header_prefix):
            para = tailored_doc.add_paragraph()

            run = para.add_run(para_text)
            run.font.name = default_font
            run.font.size = Pt(12)
            run.bold = True
        else:
            para = tailored_doc.add_paragraph()
            
            run = para.add_run(para_text)
            run.font.name = default_font
            run.font.size = default_size

        para.paragraph_format.space_after = Pt(0)  
        para.paragraph_format.line_spacing = 1.0

    # Save the new document
    tailored_doc.save(output_path)

def add_paragraph_border(paragraph):
    p = paragraph._p  # Access the XML element of the paragraph
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')  # Border type
    bottom.set(qn('w:sz'), '6')        # Border size (6 = 0.75pt)
    bottom.set(qn('w:space'), '1')     # Space between text and border
    bottom.set(qn('w:color'), 'auto')  # Border color (auto = black)

    pBdr.append(bottom)
    pPr.append(pBdr)

# Main function to run the script
def main():
    # Path to resume and job description
    resume_path = "Resume - Zeyu Shao.docx"
    job_description_path = "job_description.txt"
    output_path = "Tailored Resume - Zeyu Shao.docx"
    
    # Extract text and structure from the resume
    resume_text_and_structure = extract_text_and_structure(resume_path)
    resume_text = '\n'.join([text for text, style, run, alignment, font in resume_text_and_structure])

    # Read job description from file
    with open(job_description_path, 'r', encoding='utf-8') as file:
        job_description = file.read()

    # Tailor the resume text
    tailored_text = tailor_text(resume_text, job_description)

    # Create a new Word document with tailored content
    create_tailored_docx(tailored_text, output_path)
    
    print(f"Tailored resume saved as '{output_path}'")

if __name__ == "__main__":
    main()